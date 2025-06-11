import docx
from docx.shared import RGBColor

class FeatureExtractor:
    @staticmethod
    def extract_features(docx_path):
        doc = docx.Document(docx_path)
        features = {}

        # --- WhitespaceEncoder features (trailing spaces) ---
        trailing_spaces_per_line = []
        total_lines = 0
        for para in doc.paragraphs:
            lines = para.text.split('\n')
            for line in lines:
                total_lines += 1
                trailing = len(line) - len(line.rstrip(' '))
                trailing_spaces_per_line.append(trailing)
        features['lines_with_trailing_space'] = sum(1 for x in trailing_spaces_per_line if x > 0)
        features['max_trailing_space'] = max(trailing_spaces_per_line) if trailing_spaces_per_line else 0
        features['min_trailing_space'] = min(trailing_spaces_per_line) if trailing_spaces_per_line else 0
        features['mean_trailing_space'] = (sum(trailing_spaces_per_line) / total_lines) if total_lines else 0
        features['ratio_lines_with_trailing_space'] = (features['lines_with_trailing_space'] / total_lines) if total_lines else 0

        # --- InvisibleCharEncoder features (zero-width non-joiner) ---
        zwc = '\u200C'
        zwc_count = 0
        zwc_lines = 0
        total_chars = 0
        for para in doc.paragraphs:
            for line in para.text.split('\n'):
                total_chars += len(line)
                if zwc in line:
                    zwc_lines += 1
                    zwc_count += line.count(zwc)
        features['zero_width_char_count'] = zwc_count
        features['ratio_zero_width_char'] = (zwc_count / total_chars) if total_chars else 0
        features['lines_with_zero_width_char'] = zwc_lines

        # --- LineShiftEncoder features (line spacing) ---
        non_default_spacing = 0
        spacings = []
        for para in doc.paragraphs:
            spacing = getattr(getattr(para.paragraph_format, 'line_spacing', None), 'pt', None)
            if spacing is None:
                spacing = para.paragraph_format.line_spacing
            spacings.append(spacing)
            if spacing and abs(spacing - 1.0) > 1e-2:
                non_default_spacing += 1
        features['paragraphs_with_non_default_spacing'] = non_default_spacing
        features['ratio_paragraphs_non_default_spacing'] = (non_default_spacing / len(doc.paragraphs)) if doc.paragraphs else 0
        features['unique_line_spacings'] = len(set([s for s in spacings if s is not None]))

        # --- WordShiftEncoder features (double space after word) ---
        double_space_count = 0
        total_word_separators = 0
        double_space_paragraphs = 0
        for para in doc.paragraphs:
            words = para.text.split(' ')
            sep_count = para.text.count(' ')
            total_word_separators += sep_count
            if '  ' in para.text:
                double_space_count += para.text.count('  ')
                double_space_paragraphs += 1
        features['double_space_count'] = double_space_count
        features['ratio_double_space'] = (double_space_count / total_word_separators) if total_word_separators else 0
        features['paragraphs_with_double_space'] = double_space_paragraphs

        # --- FontVariationEncoder features (font color) ---
        color_counts = {}
        total_colored_chars = 0
        for para in doc.paragraphs:
            for run in para.runs:
                color = run.font.color.rgb if run.font.color and run.font.color.rgb else None
                if color:
                    color_tuple = (color[0], color[1], color[2])
                    color_counts[color_tuple] = color_counts.get(color_tuple, 0) + 1
                    total_colored_chars += 1
        features['font_color_variation_count'] = len(color_counts)
        features['font_color_counts'] = color_counts
        features['ratio_colored_chars'] = (total_colored_chars / total_chars) if total_chars else 0

        return features

if __name__ == "__main__":
    import sys, json
    if len(sys.argv) != 2:
        print("Usage: python feature_extractor.py <docx_file>")
    else:
        features = FeatureExtractor.extract_features(sys.argv[1])
        print(json.dumps(features, indent=2, ensure_ascii=False))
