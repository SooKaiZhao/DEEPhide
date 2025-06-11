import re
import textwrap
from docx import Document

def preprocess_text(input_text, width=70):
    """
    Cleans and splits input text into plain, readable lines.
    - Removes extra formatting (extra spaces, tabs, newlines).
    - Removes illegal XML/control characters.
    - Splits into lines of at most `width` characters.
    - Normalizes whitespace, removes control characters, lowercases text, and removes extra punctuation.
    """
    # Normalize whitespace and remove control characters
    cleaned = re.sub(r'\s+', ' ', input_text).strip()
    cleaned = re.sub(r'[^\x09\x0A\x0D\x20-\x7E\xA0-\uFFFF]', '', cleaned)
    # Lowercase for normalization (optional, can remove if case is important)
    cleaned = cleaned.lower()
    # Remove extra punctuation (optional, can remove if punctuation is important)
    cleaned = re.sub(r'["\'\-\_\*\=\+\[\]\{\}\|\\/<>]', '', cleaned)
    # Split into lines using textwrap
    lines = textwrap.wrap(cleaned, width=width)
    # Join lines with newline character
    return '\n'.join(lines)

def preprocess_docx(input_docx_path, output_docx_path, width=70):
    """
    Reads a .docx file, removes formatting, and writes each visual line as a new paragraph,
    simulating pressing 'Enter' at the end of each line (based on a fixed width).
    - Normalizes whitespace, removes control characters, lowercases text, and removes extra punctuation.
    """
    doc = Document(input_docx_path)
    new_doc = Document()
    # Only extract visible text from paragraphs (ignore tables, images, headers, footers, etc.)
    for para in doc.paragraphs:
        # Normalize whitespace and remove control characters
        cleaned = re.sub(r'\s+', ' ', para.text).strip()
        cleaned = re.sub(r'[^\x09\x0A\x0D\x20-\x7E\xA0-\uFFFF]', '', cleaned)
        cleaned = cleaned.lower()
        cleaned = re.sub(r'["\'\-\_\*\=\+\[\]\{\}\|\\/<>]', '', cleaned)
        lines = textwrap.wrap(cleaned, width=width)
        for line in lines:
            if line.strip():
                new_doc.add_paragraph(line.strip())
    # Do NOT copy tables, images, headers, footers, or any non-paragraph elements
    new_doc.save(output_docx_path)

# Example usage
if __name__ == "__main__":
    original_text = (
        "The decoding algorithm expects at least a certain number of invisible characters at the start "
        "of the file to store metadata (like the length of the hidden message). If this minimum is not "
        "met, it cannot proceed."
    )
    processed_text = preprocess_text(original_text, width=70)
    print(processed_text)
    preprocess_docx("input.docx", "output.docx", width=70)
