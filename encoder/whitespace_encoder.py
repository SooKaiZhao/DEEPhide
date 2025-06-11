from docx import Document

class WhitespaceEncoder:
    """
    Encodes a message into a cover document using trailing space steganography.
    Each line's trailing space encodes a bit: no extra space = 0, one extra space = 1.
    The message is converted to binary with a 32-bit length header.
    """
    @staticmethod
    def encode(cover_doc_path, message_path, output_path):
        try:
            # Read the message
            with open(message_path, 'r', encoding='utf-8') as f:
                message = f.read()
            message_length = len(message)
            length_binary = format(message_length, '032b')
            message_binary = ''.join(format(ord(char), '08b') for char in message)
            binary_data = length_binary + message_binary
            doc = Document(cover_doc_path)
            binary_index = 0
            total_bits = len(binary_data)
            for paragraph in doc.paragraphs:
                text = paragraph.text.rstrip(' ')
                # Only encode as many bits as we have lines
                if binary_index < total_bits:
                    bit = binary_data[binary_index]
                    if bit == '1':
                        text = text + ' '
                    # else: no extra space for '0'
                    binary_index += 1
                paragraph.clear()
                paragraph.add_run(text)
            if binary_index < total_bits:
                return False, "Not enough lines in the cover document to encode the message."
            doc.save(output_path)
            return True, "Message encoded successfully!"
        except Exception as e:
            return False, f"Encoding error: {str(e)}"
