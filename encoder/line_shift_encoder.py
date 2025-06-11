from docx import Document

class LineShiftEncoder:
    @staticmethod
    def encode(cover_doc_path, message_path, output_path):
        try:
            # Read the message
            with open(message_path, 'r', encoding='utf-8') as f:
                message = f.read()
            
            # Convert message to binary, including message length as header
            message_length = len(message)
            length_binary = format(message_length, '032b')  # 32 bits for length
            
            # Convert each character to binary and concatenate
            message_binary = ''.join(format(ord(char), '08b') for char in message)
            
            # Combine length and message
            binary_data = length_binary + message_binary
            
            # Load the document
            doc = Document(cover_doc_path)
            
            # Check if document has enough paragraphs
            if len(doc.paragraphs) < len(binary_data):
                return False, "Cover document doesn't have enough paragraphs for the message"
            
            # Apply line shift coding
            for i, paragraph in enumerate(doc.paragraphs):
                if i < len(binary_data):
                    # Adjust line spacing based on binary digit
                    if binary_data[i] == '1':
                        paragraph.paragraph_format.line_spacing = 1.15  # Shift up
                    else:
                        paragraph.paragraph_format.line_spacing = 1.0   # Normal spacing
            
            # Save the encoded document
            doc.save(output_path)
            return True, "Message encoded successfully!"
            
        except Exception as e:
            return False, f"Encoding error: {str(e)}"
