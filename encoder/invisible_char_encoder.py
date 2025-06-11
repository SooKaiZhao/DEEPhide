from docx import Document

class InvisibleCharEncoder:
    # Use zero-width non-joiner for encoding (for WPS Office compatibility)
    ZERO_WIDTH_NON_JOINER = '\u200C'  # Encodes bit '1', absence encodes '0'
    
    @staticmethod
    def encode(cover_doc_path, message_path, output_path):
        try:
            # Read the message
            with open(message_path, 'r', encoding='utf-8') as f:
                message = f.read()
            # Convert message to binary, including message length as header
            message_length = len(message)
            length_binary = format(message_length, '032b')  # 32 bits for length
            message_binary = ''.join(format(ord(char), '08b') for char in message)
            binary_data = length_binary + message_binary
            doc = Document(cover_doc_path)
            binary_index = 0
            total_bits = len(binary_data)
            for paragraph in doc.paragraphs:
                if binary_index >= total_bits:
                    break
                text = paragraph.text
                paragraph.clear()
                for char in text:
                    run = paragraph.add_run(char)
                    # Insert zero-width non-joiner for '1', nothing for '0'
                    if binary_index < total_bits:
                        if binary_data[binary_index] == '1':
                            run.add_text(InvisibleCharEncoder.ZERO_WIDTH_NON_JOINER)
                        # For '0', add nothing
                        binary_index += 1
            if binary_index < total_bits:
                return False, "Cover document doesn't have enough characters for the message"
            doc.save(output_path)
            return True, "Message encoded successfully!"
        except Exception as e:
            return False, f"Encoding error: {str(e)}"
