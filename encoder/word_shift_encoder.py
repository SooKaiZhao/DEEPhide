from docx import Document
from docx.shared import Pt

class WordShiftEncoder:
    @staticmethod
    def encode(cover_doc_path, message_path, output_path):
        try:
            # Read the message
            with open(message_path, 'r', encoding='utf-8') as f:
                message = f.read()
            
            # Convert message to binary, including message length as header
            message_length = len(message)
            length_binary = format(message_length, '032b')  # 32 bits for length
            
            # Convert each character to binary and concatenate
            message_binary = ''.join(format(ord(char), '08b') for char in message)
            
            # Combine length and message
            binary_data = length_binary + message_binary
            
            # Load the document
            doc = Document(cover_doc_path)
            
            # Count total words in document
            total_words = sum(len(paragraph.text.split()) for paragraph in doc.paragraphs)
            
            if total_words < len(binary_data):
                return False, "Cover document doesn't have enough words for the message"
            
            binary_index = 0
            
            # Apply word shift coding
            for paragraph in doc.paragraphs:
                if binary_index >= len(binary_data):
                    break
                    
                words = paragraph.text.split()
                if not words:  # Skip empty paragraphs
                    continue
                    
                # Clear existing runs and create new ones with proper spacing
                paragraph.clear()
                
                for word in words:
                    if binary_index < len(binary_data):
                        run = paragraph.add_run(word)
                        if binary_data[binary_index] == '1':
                            # Add extra space after word for '1'
                            run.add_text('  ')  # Two spaces for right shift
                        else:
                            # Normal space for '0'
                            run.add_text(' ')  # One space for left shift
                        binary_index += 1
                    else:
                        # Add remaining words with normal spacing
                        run = paragraph.add_run(word + ' ')
            
            # Save the encoded document
            doc.save(output_path)
            return True, "Message encoded successfully!"
            
        except Exception as e:
            return False, f"Encoding error: {str(e)}"
