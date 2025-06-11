from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX

class FontVariationEncoder:
    @staticmethod
    def encode(cover_doc_path, message_path, output_path):
        try:
            # Read the message
            with open(message_path, 'r', encoding='utf-8') as f:
                message = f.read()
            
            # Convert message to binary, including message length as header
            message_length = len(message)
            length_binary = format(message_length, '032b')  # 32 bits for length
            
            # Convert each character to binary and concatenate
            message_binary = ''.join(format(ord(char), '08b') for char in message)
            
            # Combine length and message
            binary_data = length_binary + message_binary
            
            # Load the document
            doc = Document(cover_doc_path)
            
            # Count total characters in document
            total_chars = sum(len(paragraph.text.replace(' ', '')) 
                            for paragraph in doc.paragraphs)
            
            if total_chars < len(binary_data):
                return False, "Cover document doesn't have enough characters for the message"
            
            binary_index = 0
            
            # Define color variations using 6-digit hex values
            color_one = RGBColor(1, 1, 1)      # Almost black (for binary 1) - #010101
            color_zero = RGBColor(2, 2, 2)     # Very slightly lighter (for binary 0) - #020202
            color_normal = RGBColor(0, 0, 0)   # Pure black (for non-encoded text) - #000000
            
            # Apply font variation coding
            for paragraph in doc.paragraphs:
                if binary_index >= len(binary_data):
                    break
                
                # Clear existing runs
                text = paragraph.text
                paragraph.clear()
                
                # Process each character
                for char in text:
                    if char.isspace():
                        # Preserve spaces without encoding
                        run = paragraph.add_run(char)
                        continue
                        
                    if binary_index < len(binary_data):
                        run = paragraph.add_run(char)
                        
                        # Base font settings
                        run.font.name = 'Calibri'
                        run.font.size = Pt(11)
                        
                        # Apply color based on binary value
                        if binary_data[binary_index] == '1':
                            run.font.color.rgb = color_one
                        else:
                            run.font.color.rgb = color_zero
                        
                        binary_index += 1
                    else:
                        # Add remaining characters without encoding
                        run = paragraph.add_run(char)
                        run.font.name = 'Calibri'
                        run.font.size = Pt(11)
                        run.font.color.rgb = color_normal
            
            # Save the encoded document
            doc.save(output_path)
            return True, "Message encoded successfully!"
            
        except Exception as e:
            return False, f"Encoding error: {str(e)}"
